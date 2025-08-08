import os
from datetime import datetime
from django.conf import settings
from docx import Document
from io import BytesIO
from django.http import HttpResponse

def substituir_termo_entrega(doc, dados):
    for par in doc.paragraphs:
        for chave, valor in dados.items():
            if chave in par.text:
                par.text = par.text.replace(chave, valor)
    return doc

def substituir_autorizacao_desconto(doc, dados):
    for par in doc.paragraphs:
        texto_completo = ''.join(run.text for run in par.runs)
        if any(chave in texto_completo for chave in dados):
            for chave in sorted(dados.keys(), key=len, reverse=True):
                texto_completo = texto_completo.replace(chave, dados[chave])

            # Remove todos os runs existentes
            for i in reversed(range(len(par.runs))):
                p_run = par.runs[i]._element
                p_run.getparent().remove(p_run)

            # Cria novo run limpo com o texto final
            par.add_run(texto_completo)
    return doc

def substituir_termo_baixa(doc, dados):
    for par in doc.paragraphs:
        for chave, valor in dados.items():
            if chave in par.text:
                par.text = par.text.replace(chave, valor)
    return doc

def substituir_termo_vinculo(doc, dados):
    def substituir_em_paragrafo(par):
        texto_completo = ''.join(run.text for run in par.runs)

        if any(chave in texto_completo for chave in dados):
            for chave in sorted(dados.keys(), key=len, reverse=True):
                texto_completo = texto_completo.replace(chave, dados[chave])

            # Remove todos os runs antigos
            for i in reversed(range(len(par.runs))):
                p_run = par.runs[i]._element
                p_run.getparent().remove(p_run)

            # Adiciona novo texto limpo
            par.add_run(texto_completo)

    # Substitui em parágrafos fora de tabelas
    for par in doc.paragraphs:
        substituir_em_paragrafo(par)

    # Substitui em parágrafos dentro de tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:
                    substituir_em_paragrafo(par)

    return doc

def gerar_documento_download(modelo_nome, nome_arquivo_download, dados):
    caminho_modelo = os.path.join(settings.BASE_DIR, 'cadastro', 'documentos_modelo', modelo_nome)
    doc = Document(caminho_modelo)

    if 'autorizacao_desconto' in modelo_nome:
        doc = substituir_autorizacao_desconto(doc, dados)
    elif 'termo_responsabilidade' in modelo_nome:
        doc = substituir_termo_entrega(doc, dados)
    elif 'termo_baixa' in modelo_nome:
        doc = substituir_termo_baixa(doc, dados)
    elif 'termo_vinculo' in modelo_nome:
        doc = substituir_termo_vinculo(doc, dados)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo_download}"'
    return response

def gerar_autorizacao_desconto(cliente, valor_total, parcelas):
    valor_parcela = valor_total / parcelas
    data_atual = datetime.now()
    dados = {
        '{{data}}': data_atual.strftime('%d/%m/%Y'),
        '{{valor}}': f"{valor_total:.2f}".replace('.', ','),
        '{{num_parcelas}}': str(parcelas),
        '{{valor_parcela}}': f"{valor_parcela:.2f}".replace('.', ','),
        '{{mês}}': data_atual.strftime('%m'),
        '{{ano}}': data_atual.strftime('%Y'),
        '{{nome}}': cliente.nome,
        '{{matricula}}': cliente.matricula,
        '{{cpf}}': cliente.cpf
    }
    nome_arquivo = f"autorizacao_desconto_os_{cliente.matricula}_{data_atual.strftime('%Y%m%d_%H%M%S')}.docx"
    return gerar_documento_download('autorizacao_desconto_modelo.docx', nome_arquivo, dados)

def gerar_termo_baixa(cliente, equipamento, motivo):
    dados = {
        '{{data}}': datetime.now().strftime('%d/%m/%Y'),
        '{{nome}}': cliente.nome,
        '{{matricula}}': cliente.matricula,
        '{{cpf}}': cliente.cpf,
        '{{modelo}}': equipamento.modelo,
        '{{numero_serie}}': equipamento.numero_serie,
        '{{responsavel_area}}': equipamento.area.responsavel_area,
        '{{motivo}}': motivo,
    }
    nome_arquivo = f"termo_baixa_{cliente.matricula}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return gerar_documento_download('termo_baixa_equipamento_modelo.docx', nome_arquivo, dados)
