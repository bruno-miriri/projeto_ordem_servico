from django.shortcuts import render, redirect, get_object_or_404
from .forms import AreaForm, ClienteForm, EquipamentoForm, OrdemServicoForm, FechamentoOSForm, BaixaEquipamentoForm
from django.utils import timezone
from .models import OrdemServico, Cliente, Equipamento, Area
import os
from django.conf import settings
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_GET
from django.db import connections, connection
from docx import Document
from datetime import datetime
from .notificacoes import enviar_email_os
from io import BytesIO
from django.views.decorators.csrf import csrf_exempt
from .documentos import substituir_autorizacao_desconto, substituir_termo_entrega, substituir_termo_baixa, substituir_termo_vinculo
from django.urls import reverse
from urllib.parse import urlencode
from django.http import HttpResponseRedirect
from django.urls import reverse


def consultar_oracle_por_matricula(empresa, matricula):
    try:
        with connections['oracle'].cursor() as cursor:
            if empresa == 'Miriri':
                cursor.execute("SELECT RA_NOME, RA_CIC FROM SRA010 WHERE RA_MAT = :matricula", {'matricula': matricula})
            elif empresa == 'Condomínio':
                cursor.execute("SELECT RA_NOME, RA_CIC FROM SRA070 WHERE RA_MAT = :matricula", {'matricula': matricula})
            else:
                return None
            return cursor.fetchone()
    except Exception:
        return None


def get_dados_cliente(request):
    empresa = request.GET.get('empresa')
    matricula = request.GET.get('matricula')

    if not empresa or not matricula:
        return JsonResponse({'sucesso': False, 'mensagem': 'Empresa ou matrícula não informada.'}, status=400)

    resultado = consultar_oracle_por_matricula(empresa, matricula)
    if resultado:
        nome, cpf = resultado
        cpf_formatado = f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"
        return JsonResponse({'sucesso': True, 'nome': nome.strip(), 'cpf_formatado': cpf_formatado})
    else:
        return JsonResponse({'sucesso': False, 'mensagem': 'Funcionário não encontrado.'}, status=404)


def menu_principal(request):
    return render(request, 'menu.html')


def cadastrar_area(request):
    if request.method == 'POST':
        form = AreaForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Área cadastrada com sucesso!')
            return redirect('cadastrar_area')
    else:
        form = AreaForm()

    areas = Area.objects.all().order_by('area')
    return render(request, 'cadastro/cadastrar_area.html', {'form': form, 'areas': areas})


def cadastrar_cliente(request):
    if request.method == 'POST':
        empresa = request.POST.get('empresa')
        matricula = request.POST.get('matricula')
        form = ClienteForm(request.POST)

        if Cliente.objects.filter(empresa=empresa, matricula=matricula).exists():
            messages.error(request, 'Cliente já está cadastrado.')
            return render(request, 'cadastro/cadastrar_cliente.html', {'form': form})

        resultado = consultar_oracle_por_matricula(empresa, matricula)
        if resultado:
            nome, cpf = resultado
            request.POST._mutable = True
            request.POST['cpf'] = cpf
            request.POST._mutable = False

        if form.is_valid():
            form.save()
            messages.success(request, 'Cliente cadastrado com sucesso!')
            return redirect('cadastrar_cliente')
        else:
            messages.error(request, 'Erro ao cadastrar cliente. Verifique os dados informados.')
    else:
        form = ClienteForm()

    return render(request, 'cadastro/cadastrar_cliente.html', {'form': form})


def verificar_cliente_existente(request):
    empresa = request.GET.get('empresa')
    matricula = request.GET.get('matricula')

    if empresa and matricula:
        existe = Cliente.objects.filter(empresa=empresa, matricula=matricula).exists()
        return JsonResponse({'existe': existe})
    return JsonResponse({'existe': False})


def cadastrar_equipamento(request):
    if request.method == 'POST':
        form = EquipamentoForm(request.POST)
        if form.is_valid():
            equipamento = form.save()

            # Redireciona com dados via GET
            base_url = reverse('cadastrar_equipamento')
            query_string = urlencode({
                'equipamento_id': equipamento.id,
                'vinculado': 1
            })
            return HttpResponseRedirect(f"{base_url}?{query_string}")
        else:
            messages.error(request, 'Erro ao cadastrar equipamento. Verifique os dados.')
    else:
        form = EquipamentoForm()

    equipamento = None
    exibe_botao_download = False

    if request.GET.get('vinculado') == '1':
        try:
            equipamento_id = int(request.GET.get('equipamento_id'))
            equipamento = Equipamento.objects.get(id=equipamento_id)
            exibe_botao_download = True
            messages.success(request, f"Equipamento '{equipamento.modelo}' cadastrado com sucesso!")
        except Exception:
            messages.warning(request, 'Equipamento cadastrado, mas não foi possível preparar o termo.')

    return render(request, 'cadastro/cadastrar_equipamento.html', {
        'form': form,
        'equipamento': equipamento,
        'exibe_botao_download': exibe_botao_download
    })



def verificar_numero_serie(request):
    numero_serie = request.GET.get('numero_serie')
    if numero_serie:
        existe = Equipamento.objects.filter(numero_serie=numero_serie).exists()
        return JsonResponse({'existe': existe})
    return JsonResponse({'existe': False})


def cadastrar_ordem_servico(request):
    if request.method == 'POST':
        form = OrdemServicoForm(request.POST, request.FILES)
        if form.is_valid():
            cliente = form.cleaned_data['cliente']
            equipamento = form.cleaned_data['equipamento']
            descricao_problema = form.cleaned_data['descricao_problema']
            valor = 0
            imagem = form.cleaned_data.get('imagem_situacao')

            imagem_path = None
            if imagem:
                nome_arquivo = f"{cliente.matricula}_{equipamento.modelo}_{timezone.now().strftime('%Y%m%d%H%M%S')}.{imagem.name.split('.')[-1]}"
                pasta_destino = os.path.join(settings.BASE_DIR, 'imagens_os')
                os.makedirs(pasta_destino, exist_ok=True)
                caminho_completo = os.path.join(pasta_destino, nome_arquivo)

                with open(caminho_completo, 'wb+') as destino:
                    for chunk in imagem.chunks():
                        destino.write(chunk)

                imagem_path = f"imagens_os/{nome_arquivo}"

            with connection.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO ordens_servico 
                        (cliente_id, equipamento_id, descricao_problema, status, valor, imagem_situacao) 
                    VALUES (%s, %s, %s, 'aberta', %s, %s)
                """, [cliente.id, equipamento.id, descricao_problema, valor, imagem_path])

                cursor.execute("SELECT LAST_INSERT_ID();")
                os_id = cursor.fetchone()[0]

            # Redireciona com parâmetros via GET
            base_url = reverse('cadastrar_ordem_servico')
            query_string = urlencode({
                'os_id': os_id,
                'cadastro': 1
            })
            return HttpResponseRedirect(f"{base_url}?{query_string}")

        else:
            messages.error(request, 'Erro ao cadastrar OS. Verifique os dados.')
    else:
        form = OrdemServicoForm()

    os_id = request.GET.get('os_id')
    exibe_botao = request.GET.get('cadastro') == '1'

    if exibe_botao and os_id:
        messages.success(request, 'Ordem de Serviço cadastrada com sucesso!')

    return render(request, 'cadastro/cadastrar_ordem_servico.html', {
        'form': form,
        'os_id': os_id,
        'exibe_botao_download': exibe_botao
    })


def download_termo_entrega(request, os_id):
    ordem = get_object_or_404(OrdemServico, id=os_id)
    cliente = ordem.cliente
    equipamento = ordem.equipamento

    resultado = consultar_oracle_por_matricula(cliente.empresa, cliente.matricula)
    cpf = resultado[1] if resultado else cliente.cpf

    dados = {
        '{{data}}': datetime.now().strftime('%d/%m/%Y'),
        '{{nome}}': cliente.nome,
        '{{matricula}}': cliente.matricula,
        '{{cpf}}': cpf,
        '{{modelo}}': equipamento.modelo,
        '{{numero_serie}}': equipamento.numero_serie,
        '{{area}}': equipamento.area.area,
        '{{responsavel_area}}': equipamento.area.responsavel_area,
    }

    nome_arquivo = f"termo_entrega_os_{cliente.matricula}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    caminho_modelo = os.path.join(settings.BASE_DIR, 'cadastro', 'documentos_modelo', 'termo_responsabilidade_modelo.docx')
    doc = Document(caminho_modelo)
    doc = substituir_termo_entrega(doc, dados)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
    return response


def fechar_ordem_servico(request):
    if request.method == 'POST':
        form = FechamentoOSForm(request.POST)
        if form.is_valid():
            ordem_servico = form.cleaned_data['os']
            valor = form.cleaned_data.get('valor')
            parcelas = form.cleaned_data.get('parcelas')

            ordem_servico.valor = valor if valor else 0
            ordem_servico.status = 'fechada'
            ordem_servico.data_fechamento = timezone.now()
            ordem_servico.save()

            messages.success(request, f'Ordem de Serviço #{ordem_servico.id} fechada com sucesso!')

            return render(request, 'cadastro/fechar_ordem_servico.html', {
                'form': FechamentoOSForm(),  # limpa o formulário
                'exibe_botao_download': bool(valor and valor > 0 and parcelas),
                'os_id_fechada': ordem_servico.id,
                'parcelas_fechadas': parcelas
            })
        else:
            messages.error(request, 'Erro ao fechar a Ordem de Serviço. Verifique os dados.')
    else:
        form = FechamentoOSForm()

    return render(request, 'cadastro/fechar_ordem_servico.html', {'form': form})

def baixar_autorizacao_desconto(request, os_id):
    ordem = get_object_or_404(OrdemServico, id=os_id)
    valor = ordem.valor
    cliente = ordem.cliente

    try:
        parcelas = int(request.GET.get('parcelas', 1))
    except (ValueError, TypeError):
        parcelas = 1

    return gerar_autorizacao_desconto(cliente, valor, parcelas)


def gerar_autorizacao_desconto(cliente, valor_total, parcelas):
    valor_parcela = valor_total / parcelas
    data_atual = datetime.now()
    dados = {
        '{{data}}': data_atual.strftime('%d/%m/%Y'),
        '{{valor}}': f"{valor_total:.2f}".replace('.', ','),
        '{{num_parcelas}}': str(parcelas),
        '{{valor_parcela}}': f"{valor_parcela:.2f}".replace('.', ','),
        '{{mês}}': data_atual.strftime('%m'),
        '{{ano}}': data_atual.strftime('%Y'),
        '{{nome}}': cliente.nome,
        '{{matricula}}': cliente.matricula,
        '{{cpf}}': cliente.cpf
    }
    nome_arquivo = f"autorizacao_desconto_os_{cliente.matricula}_{data_atual.strftime('%Y%m%d_%H%M%S')}.docx"
    caminho_modelo = os.path.join(settings.BASE_DIR, 'cadastro', 'documentos_modelo', 'autorizacao_desconto_modelo.docx')
    doc = Document(caminho_modelo)
    doc = substituir_autorizacao_desconto(doc, dados)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
    return response


def registrar_baixa_equipamento(request):
    if request.method == 'POST':
        form = BaixaEquipamentoForm(request.POST)
        if form.is_valid():
            baixa = form.save()

            # Redireciona com os dados via GET
            base_url = reverse('registrar_baixa_equipamento')
            query_string = urlencode({
                'cliente_id': baixa.cliente.id,
                'equipamento_id': baixa.equipamento.id,
                'baixado': 1
            })
            return HttpResponseRedirect(f"{base_url}?{query_string}")
        else:
            messages.error(request, 'Erro ao registrar a baixa. Verifique os dados informados.')
    else:
        form = BaixaEquipamentoForm()

    cliente = equipamento = None
    motivo = ""
    exibe_botao_download = False

    if request.GET.get('baixado') == '1':
        try:
            cliente_id = int(request.GET.get('cliente_id'))
            equipamento_id = int(request.GET.get('equipamento_id'))
            cliente = Cliente.objects.get(id=cliente_id)
            equipamento = Equipamento.objects.get(id=equipamento_id)

            # Busca o motivo mais recente dessa combinação
            from .models import BaixaEquipamento
            ultima_baixa = BaixaEquipamento.objects.filter(
                cliente=cliente,
                equipamento=equipamento
            ).order_by('-data_baixa').first()

            motivo = ultima_baixa.motivo if ultima_baixa else "Motivo não especificado"
            messages.success(request, f"Baixa registrada para o equipamento '{equipamento.modelo}' de {cliente.nome}.")
            exibe_botao_download = True
        except Exception:
            messages.warning(request, 'Baixa registrada, mas não foi possível exibir os dados para o termo.')

    return render(request, 'cadastro/registrar_baixa_equipamento.html', {
        'form': form,
        'cliente': cliente,
        'equipamento': equipamento,
        'motivo': motivo,
        'exibe_botao_download': exibe_botao_download
    })

def baixar_termo_baixa(request, cliente_id, equipamento_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    equipamento = get_object_or_404(Equipamento, id=equipamento_id)

    from .models import BaixaEquipamento
    baixa = BaixaEquipamento.objects.filter(cliente=cliente, equipamento=equipamento).order_by('-data_baixa').first()

    if not baixa:
        messages.error(request, 'Registro de baixa não encontrado.')
        return redirect('registrar_baixa_equipamento')

    return gerar_termo_baixa(cliente, equipamento, baixa.motivo)

def baixar_termo_vinculo(request, equipamento_id):
    equipamento = get_object_or_404(Equipamento, id=equipamento_id)
    cliente = equipamento.cliente  # Agora é um campo FK

    dados = {
        '{{data}}': datetime.now().strftime('%d/%m/%Y'),
        '{{nome}}': cliente.nome,
        '{{matricula}}': cliente.matricula,
        '{{cpf}}': cliente.cpf,
        '{{ugb}}': cliente.ugb,
        '{{modelo}}': equipamento.modelo,
        '{{numero_serie}}': equipamento.numero_serie,
        '{{telefone}}': cliente.telefone,
        '{{situacao}}': equipamento.situacao,
    }

    nome_arquivo = f"termo_vinculo_{cliente.matricula}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    caminho_modelo = os.path.join(settings.BASE_DIR, 'cadastro', 'documentos_modelo', 'termo_vinculo_cliente_modelo.docx')

    doc = Document(caminho_modelo)

    from .documentos import substituir_termo_vinculo
    doc = substituir_termo_vinculo(doc, dados)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
    return response




def gerar_termo_baixa(cliente, equipamento, motivo):
    dados = {
        '{{data}}': datetime.now().strftime('%d/%m/%Y'),
        '{{nome}}': cliente.nome,
        '{{matricula}}': cliente.matricula,
        '{{cpf}}': cliente.cpf,
        '{{modelo}}': equipamento.modelo,
        '{{numero_serie}}': equipamento.numero_serie,
        '{{responsavel_area}}': equipamento.area.responsavel_area,
        '{{motivo}}': motivo,
    }
    nome_arquivo = f"termo_baixa_{cliente.matricula}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    caminho_modelo = os.path.join(settings.BASE_DIR, 'cadastro', 'documentos_modelo', 'termo_baixa_equipamento_modelo.docx')
    doc = Document(caminho_modelo)
    doc = substituir_termo_baixa(doc, dados)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
    return response